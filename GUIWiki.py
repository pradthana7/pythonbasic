# GUIWiki.py
import wikipedia

# python to docx
from docx import Document
def Wiki(keyword,lang='th'):
    wikipedia.set_lang(lang)
    # summary สำหรับบทความที่่สรป
    data = wikipedia.summary(keyword)

    # page + content บทความทั้งหน้า
    data2 = wikipedia.page(keyword)
    data2 = data2.content


    doc = Document() #สร้างไฟล์ word ในpython
    doc.add_heading(keyword,0)

    doc.add_paragraph(data2)
    doc.save(keyword + '.docx')
    print('สร้างไฟล์สำเร็จ')

#เปลี่ยนเป็นภาษาไทย
wikipedia.set_lang('th')
from tkinter import *
from tkinter import ttk
from tkinter import messagebox


GUI = Tk() # ต้องคู่กับ GUI.mainloop
GUI.title('โปรแกรม wiki')
GUI.geometry('400x300')
# config
FONT1 = ('Angsana new',15)

# คำอธิบาย
L = ttk.Label(GUI, text='ค้นหาบทความ',font=FONT1)
L.pack()

# ช่องค้นหาข้อมูล
v_search = StringVar() # กล่องสำหรับเก็บkeyword
E1 = ttk.Entry(GUI,textvariable = v_search,font=FONT1,width=40)
E1.pack(pady=10)

# ปุ่มค้นหา
def Search():
    keyword = v_search.get() # .get() คือดึงข้อมูลเข้้ามา
    try:
        # ลองค้นหาดูว่าได้ผลลัพธ์หรือไม่ หาได้ให้ผ่านไป
        language = v_radio.get() # th / en / zh
        Wiki(keyword,language)
        messagebox.showinfo('บันทึกสำเร็จ','ค้นหาข้อความสำเร็จ บันทึกเรียบร้อยแล้ว')
    except:
        # หารันคำสั่งแล้วมีปัญหา แสดงข้อความแจ้งเืตือน
        messagebox.showwarning('Keyword Error','กรุณากรอกคำค้นหาใหม่')
        
        
    # print(wikipedia.search(keyword))
    # result = wikipedia.summary(keyword)
    # print(result)

    
B1 = ttk.Button(GUI,text='Search',command=Search)
B1.pack(ipadx=20,ipady=10) #ipadx ขยายภาพแกนx,ipady ขยายภาพแกนy

# เลือกภาษา
F1 = Frame(GUI)
F1.pack(pady=10)

v_radio = StringVar() # ช่องเก็บข้อมูลภาษา #Radiobutton ทำได้เฉพาะคำเดียวหรือตัวเดียว

RB1 = ttk.Radiobutton(F1,text='ภาษาไทย',variable=v_radio,value='th')
RB2 = ttk.Radiobutton(F1,text='อังกฤษ',variable=v_radio,value='en')
RB3 = ttk.Radiobutton(F1,text='จีน',variable=v_radio,value='zh')
RB1.invoke() # สั่งให้ค่าเริ่มต้นเป็นภาษาไทย

RB1.grid(row=0,column=0)
RB2.grid(row=0,column=1)
RB3.grid(row=0,column=2)

GUI.mainloop()

