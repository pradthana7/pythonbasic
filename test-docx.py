# test-docx.py
from docx import Document
import wikipedia
wikipedia.set_lang('th')

data = wikipedia.summary('แมว')

doc = Document() #สร้างไฟล์ word ในpython
doc.add_heading('แมว',0)

doc.add_paragraph(data)
doc.save('แมว.docx')
print('สร้างไฟล์สำเร็จ')


