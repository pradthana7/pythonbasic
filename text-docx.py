# test-docx.py
from docx import Document
import wikipedia


def wiki(keyword,lang='th'):
    wikipedia.set_lang(lang)
    # summary สำหรับบทความที่่สรุป
    data = wikipedia.summary(keyword)

    # page + content บทความทั้งหน้า
    data2 = wikipedia.page(keyword)
    data2 = data2.content


    doc = Document() #สร้างไฟล์ word ในpython
    doc.add_heading(keyword,0)

    doc.add_paragraph(data2)
    doc.save(keyword + '.docx')
    print('สร้างไฟล์สำเร็จ')

try:
    wiki('')
wiki('sdiiwwdfflrwwp','zh')
except:
    print('error')

#   wiki('united state of america','en')
#   wiki('ประเทศญี่ปุ่น','jp')


